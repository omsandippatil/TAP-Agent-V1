import io
import os
import datetime
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

from app.pipeline.llm import LLM_UNAVAILABLE_EVIDENCE

FONT_NAME = "Calibri"

TEAL_DARK = RGBColor(0x0F, 0x3D, 0x3E)
TEAL = RGBColor(0x0F, 0x76, 0x6E)
TEAL_MID = RGBColor(0x14, 0x6B, 0x65)
TEAL_LIGHT = RGBColor(0x20, 0xB2, 0xAA)
YELLOW = RGBColor(0xF5, 0xC5, 0x18)
YELLOW_DARK = RGBColor(0x8A, 0x62, 0x00)
INK = RGBColor(0x1F, 0x29, 0x37)
GREY = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
AMBER = RGBColor(0xD9, 0x77, 0x06)
RED = RGBColor(0xDC, 0x26, 0x26)
LINKEDIN_BLUE = RGBColor(0x0A, 0x66, 0xC2)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CREAM_HEX = "FAF7F2"
TEAL_DARK_HEX = "0F3D3E"
TEAL_MID_HEX = "146B65"
TEAL_SOFT_HEX = "E6F5F3"
YELLOW_HEX = "F5C518"
YELLOW_SOFT_HEX = "FEF7DC"

LOGO_PATH = os.path.join(os.path.dirname(__file__), "assets", "ff_logo.png")

SOURCE_LABELS = {
    "india_csr_page": "Company India CSR page",
    "mca_portal": "MCA portal (verified)",
    "mca_via_search": "MCA via web search (proxy — not a verified filing)",
    "national_csr_portal": "National CSR Portal (csr.gov.in)",
    "annual_report": "Annual / sustainability report",
    "global_annual_report": "Annual / sustainability report",
    "partner_search": "Partner-focused web search",
    "people_search": "LinkedIn / people search snippets",
    "plans_search": "Partnerships & announced-plans search",
    "sector_eligibility_search": "Sector & eligibility search",
}


def fit_hex(score):
    if score >= 90:
        return TEAL_DARK_HEX
    if score >= 80:
        return TEAL_MID_HEX
    if score >= 65:
        return "0EA5A0"
    if score >= 45:
        return "D97706"
    return "DC2626"


def fit_color(score):
    if score >= 90:
        return TEAL_DARK
    if score >= 80:
        return TEAL_MID
    if score >= 65:
        return TEAL_LIGHT
    if score >= 45:
        return AMBER
    return RED


def fit_label(score):
    if score >= 90:
        return "PRIORITY MATCH"
    if score >= 80:
        return "STRONG MATCH"
    if score >= 65:
        return "PROMISING"
    if score >= 45:
        return "WATCHLIST"
    return "LOW FIT"


TC_PR_CHILD_ORDER = ("cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "cellIns", "cellDel", "cellMerge", "tcPrChange")


def insert_tc_pr_child(tc_pr, new_element, tag_name):
    insert_index = TC_PR_CHILD_ORDER.index(tag_name)
    for existing in tc_pr:
        existing_tag = existing.tag.split("}")[-1]
        if existing_tag in TC_PR_CHILD_ORDER and TC_PR_CHILD_ORDER.index(existing_tag) > insert_index:
            existing.addprevious(new_element)
            return
    tc_pr.append(new_element)


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_color)
    insert_tc_pr_child(tc_pr, shading, "shd")


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    insert_tc_pr_child(tc_pr, margins, "tcMar")


TBL_PR_CHILD_ORDER = ("tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout", "tblCellMar", "tblLook")


def insert_tbl_pr_child(tbl_pr, new_element, tag_name):
    insert_index = TBL_PR_CHILD_ORDER.index(tag_name)
    for existing in tbl_pr:
        existing_tag = existing.tag.split("}")[-1]
        if existing_tag in TBL_PR_CHILD_ORDER and TBL_PR_CHILD_ORDER.index(existing_tag) > insert_index:
            existing.addprevious(new_element)
            return
    tbl_pr.append(new_element)


def set_table_borders(table, color="E1E9E7", size=4):
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    insert_tbl_pr_child(tbl_pr, borders, "tblBorders")


def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    insert_tbl_pr_child(tbl_pr, borders, "tblBorders")


def add_hyperlink(paragraph, url, display_text, color=TEAL_MID_HEX, bold=False):
    part = paragraph.part
    relationship_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    font_element = OxmlElement("w:rFonts")
    font_element.set(qn("w:ascii"), FONT_NAME)
    font_element.set(qn("w:hAnsi"), FONT_NAME)
    run_properties.append(font_element)
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    run_properties.append(color_element)
    underline_element = OxmlElement("w:u")
    underline_element.set(qn("w:val"), "single")
    run_properties.append(underline_element)
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), "18")
    run_properties.append(size_element)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = display_text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


PPR_CHILD_ORDER = ("pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection", "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr", "sectPr", "pPrChange")


def insert_ppr_child(ppr, new_element, tag_name):
    insert_index = PPR_CHILD_ORDER.index(tag_name)
    for existing in ppr:
        existing_tag = existing.tag.split("}")[-1]
        if existing_tag in PPR_CHILD_ORDER and PPR_CHILD_ORDER.index(existing_tag) > insert_index:
            existing.addprevious(new_element)
            return
    ppr.append(new_element)


def add_section_heading(doc, text, size=13, color=TEAL_DARK, space_before=16):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    border_element = OxmlElement("w:pBdr")
    bottom_element = OxmlElement("w:bottom")
    bottom_element.set(qn("w:val"), "single")
    bottom_element.set(qn("w:sz"), "6")
    bottom_element.set(qn("w:space"), "3")
    bottom_element.set(qn("w:color"), TEAL_SOFT_HEX)
    border_element.append(bottom_element)
    insert_ppr_child(paragraph._p.get_or_add_pPr(), border_element, "pBdr")
    run = paragraph.add_run(text.upper())
    run.font.name = FONT_NAME
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return paragraph


def style_small(run_or_paragraph, size=8.5, color=GREY, italic=True):
    font = run_or_paragraph.font
    font.name = FONT_NAME
    font.size = Pt(size)
    font.color.rgb = color
    font.italic = italic


def add_evidence_paragraph(doc, excerpt_text, url=""):
    if not excerpt_text:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"\u201c{excerpt_text[:260]}\u2026\u201d  ")
    run.font.name = FONT_NAME
    style_small(run)
    if url:
        add_hyperlink(paragraph, url, "source")


_HIGHLIGHT_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def add_marked_run(paragraph, text, base_size=9.5, base_bold=False, base_color=None):
    if not text:
        return
    position = 0
    for match in _HIGHLIGHT_PATTERN.finditer(text):
        if match.start() > position:
            plain_run = paragraph.add_run(text[position:match.start()])
            plain_run.font.name = FONT_NAME
            plain_run.font.size = Pt(base_size)
            plain_run.bold = base_bold
            if base_color:
                plain_run.font.color.rgb = base_color
        highlight_run = paragraph.add_run(match.group(1))
        highlight_run.font.name = FONT_NAME
        highlight_run.font.size = Pt(base_size)
        highlight_run.bold = True
        highlight_run.font.color.rgb = YELLOW_DARK
        highlight_run.font.highlight_color = 7
        position = match.end()
    if position < len(text):
        tail_run = paragraph.add_run(text[position:])
        tail_run.font.name = FONT_NAME
        tail_run.font.size = Pt(base_size)
        tail_run.bold = base_bold
        if base_color:
            tail_run.font.color.rgb = base_color


def add_marked_paragraph(doc, text, base_size=9.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    add_marked_run(paragraph, text, base_size=base_size)
    return paragraph


def add_source_reference_paragraph(doc, source_name, source_lookup, base_size=8.5):
    if not source_name:
        return
    url = source_lookup.get(source_name, "")
    label = SOURCE_LABELS.get(source_name, source_name)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.space_after = Pt(4)
    prefix_run = paragraph.add_run("Source: ")
    prefix_run.font.name = FONT_NAME
    style_small(prefix_run, base_size)
    if url:
        add_hyperlink(paragraph, url, label, color="6B7280")
    else:
        label_run = paragraph.add_run(label)
        style_small(label_run, base_size)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)


def style_header_cell(cell, text, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = WHITE
    shade_cell(cell, TEAL_MID_HEX)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_body_cell(cell, text, size=8.5, bold=False, color=INK, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    if align:
        paragraph.alignment = align
    run = paragraph.add_run(str(text) if text is not None else "")
    run.font.name = FONT_NAME
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return paragraph


def style_link_cell(cell, url, display_text="LinkedIn", size=8.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if url:
        add_hyperlink(paragraph, url, display_text, color="0A66C2", bold=True)
    else:
        run = paragraph.add_run("—")
        run.font.name = FONT_NAME
        run.font.size = Pt(size)
        run.font.color.rgb = GREY
    return paragraph


def build_data_table(doc, headers, col_widths_cm, rows_data, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    for idx, heading_text in enumerate(headers):
        style_header_cell(table.rows[0].cells[idx], heading_text)
    for row_index, values in enumerate(rows_data, start=1):
        row_cells = table.add_row().cells
        for col_index, value in enumerate(values):
            style_body_cell(row_cells[col_index], value)
        if zebra and row_index % 2 == 0:
            for cell in row_cells:
                shade_cell(cell, TEAL_SOFT_HEX)
    set_col_widths(table, col_widths_cm)
    return table


def merge_decision_makers(result: dict) -> list:
    analysis = result.get("analysis") or {}
    llm_people = analysis.get("decision_makers") or []
    scraped_people = result.get("decision_makers") or []

    scraped_by_name = {}
    for person in scraped_people:
        key = (person.get("name") or "").strip().lower()
        if key and key not in scraped_by_name:
            scraped_by_name[key] = person

    merged = []
    seen_names = set()

    for person in llm_people:
        name = (person.get("name") or "").strip()
        if not name:
            continue
        key = name.lower()
        seen_names.add(key)
        scraped_match = scraped_by_name.get(key, {})
        linkedin_url = person.get("linkedin_url") or scraped_match.get("url", "")
        merged.append({
            "name": name,
            "title": person.get("title") or scraped_match.get("title", ""),
            "linkedin_url": linkedin_url,
            "tenure_status": person.get("tenure_status", "UNKNOWN"),
            "source_excerpt": person.get("source_excerpt", ""),
            "source": person.get("source", ""),
        })

    for person in scraped_people:
        name = (person.get("name") or "").strip()
        key = name.lower()
        if not name or key in seen_names:
            continue
        seen_names.add(key)
        merged.append({
            "name": name,
            "title": person.get("title", ""),
            "linkedin_url": person.get("url", ""),
            "tenure_status": "UNKNOWN",
            "source_excerpt": "",
            "source": "people_search",
        })

    return merged


async def generate_docx_report(company: str, result: dict, mode: str = "deep") -> bytes:
    fit_score = result.get("fit_score", 0)
    state = result.get("state", "")
    insight = result.get("strategic_insight", "")
    analysis = result.get("analysis") or {}
    breakdown = result.get("score_breakdown", {}) or {}
    sources = result.get("sources", []) or []
    source_lookup = {s.get("source_name", ""): s.get("url", "") for s in sources if s.get("url")}
    decision_makers = merge_decision_makers(result)
    important_links = result.get("important_links", []) or []
    tier = result.get("scoring_tier", {}) or {}
    has_analysis = bool(analysis)
    generated_on = datetime.datetime.now().strftime("%d %B %Y")

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.6)
        section.left_margin = section.right_margin = Cm(1.8)
    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = INK
    normal_style.paragraph_format.space_after = Pt(4)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header_table = doc.add_table(rows=1, cols=2 if os.path.exists(LOGO_PATH) else 1)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_table_borders(header_table)

    if os.path.exists(LOGO_PATH):
        logo_cell = header_table.rows[0].cells[0]
        shade_cell(logo_cell, TEAL_DARK_HEX)
        set_cell_margins(logo_cell, top=160, bottom=160, left=160, right=100)
        logo_cell.width = Cm(2.1)
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(LOGO_PATH, width=Cm(1.5), height=Cm(1.5))
        text_cell = header_table.rows[0].cells[1]
        text_cell.width = Cm(14.5)
    else:
        text_cell = header_table.rows[0].cells[0]

    shade_cell(text_cell, TEAL_DARK_HEX)
    set_cell_margins(text_cell, top=160, bottom=160, left=180, right=180)
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header_paragraph = text_cell.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.space_after = Pt(2)
    header_run = header_paragraph.add_run("THE APPRENTICE PROJECT — CSR RESEARCH BRIEF")
    header_run.font.name = FONT_NAME
    header_run.bold = True
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = YELLOW
    company_paragraph = text_cell.add_paragraph()
    company_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_paragraph.paragraph_format.space_after = Pt(2)
    company_run = company_paragraph.add_run(company)
    company_run.font.name = FONT_NAME
    company_run.bold = True
    company_run.font.size = Pt(20)
    company_run.font.color.rgb = WHITE
    subtitle_paragraph = text_cell.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.paragraph_format.space_after = Pt(0)
    subtitle_run = subtitle_paragraph.add_run(
        f"{'Deep Research' if mode == 'deep' else 'Prospect Screening'}  ·  {generated_on}  ·  "
        f"Prepared by TAP Fundraising Research Agent")
    subtitle_run.font.name = FONT_NAME
    subtitle_run.font.size = Pt(8.5)
    subtitle_run.font.color.rgb = RGBColor(0xB8, 0xE3, 0xE0)

    add_section_heading(doc, "Executive summary")
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    no_table_borders(summary_table)
    set_col_widths(summary_table, [4.2, 12.4])
    score_cell = summary_table.rows[0].cells[0]
    shade_cell(score_cell, fit_hex(fit_score))
    set_cell_margins(score_cell, top=200, bottom=200, left=100, right=100)
    score_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    score_paragraph = score_cell.paragraphs[0]
    score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_paragraph.paragraph_format.space_after = Pt(0)
    score_run = score_paragraph.add_run(f"{fit_score}")
    score_run.font.name = FONT_NAME
    score_run.bold = True
    score_run.font.size = Pt(34)
    score_run.font.color.rgb = WHITE
    label_paragraph = score_cell.add_paragraph()
    label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_paragraph.paragraph_format.space_after = Pt(0)
    label_run = label_paragraph.add_run(f"{fit_label(fit_score)} / 100")
    label_run.font.name = FONT_NAME
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = WHITE

    insight_cell = summary_table.rows[0].cells[1]
    shade_cell(insight_cell, CREAM_HEX)
    set_cell_margins(insight_cell, top=160, bottom=160, left=220, right=180)
    insight_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    insight_paragraph = insight_cell.paragraphs[0]
    insight_paragraph.paragraph_format.space_after = Pt(6)
    add_marked_run(insight_paragraph, insight, base_size=9.5, base_color=INK)
    state_labels = {"FOUND": "Evidence found across sources",
                     "NOT_FOUND_IN_SOURCE": "Partial research — some sources empty",
                     "CONFIRMED_ABSENT": "No public India CSR evidence"}
    state_paragraph = insight_cell.add_paragraph()
    state_paragraph.paragraph_format.space_after = Pt(0)
    state_run = state_paragraph.add_run(f"Research state: {state_labels.get(state, state)}")
    style_small(state_run, 8, GREY, italic=False)

    add_section_heading(doc, "CSR delivery model")
    delivery_model = analysis.get("delivery_model", "UNCLEAR") if has_analysis else "UNCLEAR"
    delivery_color = {"FUNDER": TEAL_MID, "HYBRID": TEAL_DARK, "IMPLEMENTER": AMBER}.get(delivery_model, GREY)
    model_paragraph = doc.add_paragraph()
    model_paragraph.paragraph_format.space_after = Pt(6)
    model_run = model_paragraph.add_run(delivery_model + "   ")
    model_run.font.name = FONT_NAME
    model_run.bold = True
    model_run.font.size = Pt(11)
    model_run.font.color.rgb = delivery_color
    if has_analysis:
        add_marked_run(model_paragraph, analysis.get("delivery_model_evidence", ""), base_size=9.5)
    if has_analysis and analysis.get("delivery_model_source"):
        add_source_reference_paragraph(doc, analysis["delivery_model_source"], source_lookup)

    add_section_heading(doc, "Semantic alignment & authenticity")
    if has_analysis:
        metrics_paragraph = doc.add_paragraph()
        metrics_paragraph.paragraph_format.space_after = Pt(6)
        metrics_run = metrics_paragraph.add_run(
            f"Semantic alignment: {analysis.get('overall_semantic_alignment', 0)}/100    ·    "
            f"Evidence authenticity: {analysis.get('overall_authenticity_score', 0)}/100    ·    "
            f"Avg criteria confidence: {breakdown.get('average_confidence_pct', 0)}%"
        )
        metrics_run.font.name = FONT_NAME
        metrics_run.bold = True
        metrics_run.font.size = Pt(10)
        metrics_run.font.color.rgb = TEAL_DARK
        add_marked_paragraph(doc, analysis.get("alignment_rationale", ""))
    else:
        unavailable_run = doc.add_paragraph().add_run(LLM_UNAVAILABLE_EVIDENCE)
        unavailable_run.font.name = FONT_NAME
        unavailable_run.font.color.rgb = RED

    add_section_heading(doc, "Fit rationale")
    if has_analysis:
        add_marked_paragraph(doc, analysis.get("fit_rationale", ""), base_size=9.5)
    else:
        unavailable_run = doc.add_paragraph().add_run(LLM_UNAVAILABLE_EVIDENCE)
        unavailable_run.font.name = FONT_NAME
        unavailable_run.font.color.rgb = RED

    add_section_heading(doc, "Criteria scorecard (0–5)")
    if has_analysis:
        criteria = analysis.get("criteria", [])
        criteria_table = doc.add_table(rows=1, cols=4)
        criteria_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_borders(criteria_table)
        for index, heading_text in enumerate(("Criterion", "Score", "Confidence", "Evidence")):
            style_header_cell(criteria_table.rows[0].cells[index], heading_text)
        for row_index, criterion in enumerate(criteria, start=1):
            row_cells = criteria_table.add_row().cells
            style_body_cell(row_cells[0], criterion["name"], bold=True)
            style_body_cell(row_cells[1], criterion["score"], align=WD_ALIGN_PARAGRAPH.CENTER)
            style_body_cell(row_cells[2], f"{criterion.get('confidence', 0)}%", align=WD_ALIGN_PARAGRAPH.CENTER)
            row_cells[3].text = ""
            evidence_paragraph = row_cells[3].paragraphs[0]
            evidence_paragraph.paragraph_format.space_after = Pt(0)
            set_cell_margins(row_cells[3])
            row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            add_marked_run(evidence_paragraph, criterion.get("evidence", ""), base_size=8.5)
            criterion_source = criterion.get("source", "")
            if criterion_source:
                source_url = source_lookup.get(criterion_source, "")
                if source_url:
                    footer_run = evidence_paragraph.add_run("  ")
                    footer_run.font.size = Pt(7.5)
                    add_hyperlink(evidence_paragraph, source_url,
                                  f"[{SOURCE_LABELS.get(criterion_source, criterion_source)}]", color="6B7280")
            if row_index % 2 == 0:
                for cell in row_cells:
                    shade_cell(cell, TEAL_SOFT_HEX)
            if criterion.get("confidence", 100) < 50:
                shade_cell(row_cells[2], YELLOW_SOFT_HEX)
        set_col_widths(criteria_table, [4.6, 1.7, 2.2, 8.0])
        note_paragraph = doc.add_paragraph()
        note_paragraph.paragraph_format.space_before = Pt(8)
        note_run = note_paragraph.add_run(
            "Every score and evidence line is generated by the LLM from the fetched evidence "
            "text for this specific company. Low-confidence rows are highlighted — a person "
            "verifies every figure before use.")
        style_small(note_run)
        if analysis.get("open_questions"):
            questions_paragraph = doc.add_paragraph()
            questions_paragraph.paragraph_format.space_after = Pt(0)
            questions_run = questions_paragraph.add_run(
                "Open questions: " + " · ".join(analysis["open_questions"][:5]))
            style_small(questions_run, 8.5, GREY, italic=False)
    else:
        unavailable_paragraph = doc.add_paragraph()
        unavailable_run = unavailable_paragraph.add_run(LLM_UNAVAILABLE_EVIDENCE)
        unavailable_run.font.name = FONT_NAME
        unavailable_run.font.color.rgb = RED

    if has_analysis and analysis.get("red_flags"):
        add_section_heading(doc, "Red flags")
        for flag in analysis["red_flags"]:
            flag_paragraph = doc.add_paragraph()
            flag_paragraph.paragraph_format.space_after = Pt(6)
            severity_color = RED if flag.get("severity") == "high" else AMBER
            flag_run = flag_paragraph.add_run(f"{flag.get('flag', '')} ({flag.get('severity', '')})   ")
            flag_run.font.name = FONT_NAME
            flag_run.bold = True
            flag_run.font.color.rgb = severity_color
            flag_run.font.size = Pt(9.5)
            explanation_run = flag_paragraph.add_run(flag.get("explanation", ""))
            explanation_run.font.name = FONT_NAME
            explanation_run.font.size = Pt(9.5)
            if flag.get("source"):
                add_source_reference_paragraph(doc, flag["source"], source_lookup)

    spend = analysis.get("spend", {}) if has_analysis else {}
    add_section_heading(doc, "India CSR spend")
    spend_paragraph = doc.add_paragraph()
    spend_paragraph.paragraph_format.space_after = Pt(4)
    spend_run = spend_paragraph.add_run(
        f"{spend.get('display', 'Not publicly disclosed')}"
        + (f"   ({spend['fiscal_year']})" if spend.get("fiscal_year") else ""))
    spend_run.font.name = FONT_NAME
    spend_run.bold = True
    spend_run.font.size = Pt(11)
    spend_run.font.color.rgb = TEAL_DARK
    if spend.get("source_excerpt"):
        add_evidence_paragraph(doc, spend["source_excerpt"], source_lookup.get(spend.get("source", ""), ""))
    if spend.get("trend_direction") and spend.get("trend_direction") != "UNKNOWN":
        trend_paragraph = doc.add_paragraph()
        trend_paragraph.paragraph_format.space_after = Pt(4)
        trend_run = trend_paragraph.add_run(f"Trend: {spend['trend_direction']}   ")
        trend_run.font.name = FONT_NAME
        trend_run.bold = True
        trend_run.font.size = Pt(9.5)
        trend_run.font.color.rgb = TEAL_MID
        add_marked_run(trend_paragraph, spend.get("trend_evidence", ""), base_size=9)
    if spend.get("history"):
        history_rows = [
            (entry.get("fiscal_year", ""), entry.get("display", ""), entry.get("source_excerpt", ""))
            for entry in spend["history"]
        ]
        build_data_table(doc, ["Fiscal year", "Amount", "Excerpt"], [3.2, 3.2, 9.9], history_rows)

    add_section_heading(doc, "Programmes & initiatives")
    programmes = analysis.get("programmes", []) if has_analysis else []
    if programmes:
        for programme in programmes:
            programme_paragraph = doc.add_paragraph(style="List Bullet")
            programme_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            programme_paragraph.paragraph_format.space_after = Pt(2)
            name_run = programme_paragraph.add_run(programme.get("name", ""))
            name_run.font.name = FONT_NAME
            name_run.bold = True
            name_run.font.size = Pt(9.5)
            name_run.font.color.rgb = TEAL_DARK
            if programme.get("is_multi_year"):
                multi_year_run = programme_paragraph.add_run("  MULTI-YEAR")
                multi_year_run.font.name = FONT_NAME
                multi_year_run.font.size = Pt(7.5)
                multi_year_run.bold = True
                multi_year_run.font.color.rgb = YELLOW_DARK
                multi_year_run.font.highlight_color = 7
            if programme.get("cohort_or_scale"):
                scale_run = programme_paragraph.add_run(f"   · {programme['cohort_or_scale']}")
                scale_run.font.name = FONT_NAME
                scale_run.font.size = Pt(8.5)
            if programme.get("description"):
                add_evidence_paragraph(doc, programme["description"], source_lookup.get(programme.get("source", ""), ""))
    else:
        note_run = doc.add_paragraph().add_run("No named programmes found in fetched sources.")
        note_run.font.name = FONT_NAME
        style_small(note_run)

    add_section_heading(doc, "Funded / implementation partners")
    partners = analysis.get("partners", []) if has_analysis else []
    if partners:
        partner_rows = [
            (partner.get("name", ""), partner.get("relationship_type", ""), partner.get("source_excerpt", ""))
            for partner in partners
        ]
        build_data_table(doc, ["Partner", "Type", "Evidence"], [4.6, 3.2, 8.5], partner_rows)
        note_paragraph = doc.add_paragraph()
        note_paragraph.paragraph_format.space_before = Pt(8)
        note_run = note_paragraph.add_run(
            "Partners are listed only when named in a fetched source for this company, in CSR context.")
        style_small(note_run)
    else:
        note_run = doc.add_paragraph().add_run("No funded partners found in public sources.")
        note_run.font.name = FONT_NAME
        style_small(note_run)

    add_section_heading(doc, "CSR decision-makers")
    if decision_makers:
        people_table = doc.add_table(rows=1, cols=5)
        people_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_borders(people_table)
        for index, heading_text in enumerate(["Name", "Title / role", "Tenure", "LinkedIn", "Evidence"]):
            style_header_cell(people_table.rows[0].cells[index], heading_text)
        for row_index, person in enumerate(decision_makers, start=1):
            row_cells = people_table.add_row().cells
            style_body_cell(row_cells[0], person.get("name", ""), bold=True)
            style_body_cell(row_cells[1], person.get("title", ""))
            tenure_display = str(person.get("tenure_status", "UNKNOWN")).replace("_", " ").title()
            is_new = person.get("tenure_status") == "NEW_UNDER_1YR"
            style_body_cell(row_cells[2], tenure_display, bold=is_new,
                             color=YELLOW_DARK if is_new else INK, align=WD_ALIGN_PARAGRAPH.CENTER)
            style_link_cell(row_cells[3], person.get("linkedin_url", ""))
            style_body_cell(row_cells[4], person.get("source_excerpt", ""), size=8.5)
            if is_new:
                shade_cell(row_cells[2], YELLOW_SOFT_HEX)
            elif row_index % 2 == 0:
                for cell in row_cells:
                    shade_cell(cell, TEAL_SOFT_HEX)
        set_col_widths(people_table, [3.4, 3.8, 2.0, 2.1, 5.0])
        note_paragraph = doc.add_paragraph()
        note_paragraph.paragraph_format.space_before = Pt(8)
        note_run = note_paragraph.add_run(
            "Only decision-makers named in a fetched source for this company are shown — always "
            "verify current title and contact details before outreach.")
        style_small(note_run)
    else:
        note_run = doc.add_paragraph().add_run("No CSR decision-makers found in public sources.")
        note_run.font.name = FONT_NAME
        style_small(note_run)

    add_section_heading(doc, "Governance & commitment signals")
    board_affinity = analysis.get("board_affinity", {}) if has_analysis else {}
    volunteering = analysis.get("volunteering", {}) if has_analysis else {}
    governance_paragraph = doc.add_paragraph()
    governance_paragraph.paragraph_format.space_after = Pt(4)
    governance_run = governance_paragraph.add_run("Board / promoter education affinity:   ")
    governance_run.font.name = FONT_NAME
    governance_run.bold = True
    governance_run.font.size = Pt(9.5)
    if board_affinity.get("present"):
        add_marked_run(
            governance_paragraph,
            f"{board_affinity.get('person_name', '')} — {board_affinity.get('connection', '')}",
            base_size=9.5,
        )
    else:
        no_run = governance_paragraph.add_run("No named board or promoter affinity found.")
        style_small(no_run, 9, GREY)
    if board_affinity.get("present") and board_affinity.get("source"):
        add_source_reference_paragraph(doc, board_affinity["source"], source_lookup)

    volunteering_paragraph = doc.add_paragraph()
    volunteering_paragraph.paragraph_format.space_after = Pt(4)
    volunteering_run = volunteering_paragraph.add_run("Employee volunteering / payroll giving:   ")
    volunteering_run.font.name = FONT_NAME
    volunteering_run.bold = True
    volunteering_run.font.size = Pt(9.5)
    if volunteering.get("present"):
        add_marked_run(
            volunteering_paragraph,
            f"{volunteering.get('programme_name', '')} — {volunteering.get('description', '')}",
            base_size=9.5,
        )
    else:
        no_run = volunteering_paragraph.add_run("No named volunteering programme found.")
        style_small(no_run, 9, GREY)
    if volunteering.get("present") and volunteering.get("source"):
        add_source_reference_paragraph(doc, volunteering["source"], source_lookup)

    eligibility = analysis.get("eligibility", {}) if has_analysis else {}
    group_foundation = analysis.get("group_foundation", {}) if has_analysis else {}
    add_section_heading(doc, "Sector & eligibility")
    sector = analysis.get("sector", {}) if has_analysis else {}
    if sector.get("sector") and sector.get("sector") != "UNKNOWN":
        sector_paragraph = doc.add_paragraph()
        sector_paragraph.paragraph_format.space_after = Pt(4)
        sector_run = sector_paragraph.add_run(
            sector["sector"] + (f"   ·   {sector['sub_sector']}" if sector.get("sub_sector") else ""))
        sector_run.font.name = FONT_NAME
        sector_run.bold = True
        sector_run.font.size = Pt(10.5)
        sector_run.font.color.rgb = TEAL_DARK
        if sector.get("reasoning"):
            add_marked_paragraph(doc, sector["reasoning"], base_size=9)
    else:
        note_run = doc.add_paragraph().add_run("Sector not confidently identified from fetched sources.")
        note_run.font.name = FONT_NAME
        style_small(note_run, 9, GREY)
    eligibility_paragraph = doc.add_paragraph()
    eligibility_paragraph.paragraph_format.space_after = Pt(4)
    eligibility_run = eligibility_paragraph.add_run(
        f"Section 135 mandate: {eligibility.get('plausibly_mandated', 'UNKNOWN')}")
    eligibility_run.font.name = FONT_NAME
    eligibility_run.bold = True
    eligibility_run.font.size = Pt(9.5)
    eligibility_run.font.color.rgb = (
        GREEN if eligibility.get("plausibly_mandated") == "LIKELY"
        else AMBER if eligibility.get("plausibly_mandated") == "UNLIKELY" else GREY
    )
    if eligibility.get("reasoning"):
        add_evidence_paragraph(doc, eligibility["reasoning"], source_lookup.get(eligibility.get("source", ""), ""))
    if group_foundation.get("routed_through_group"):
        group_paragraph = doc.add_paragraph()
        group_paragraph.paragraph_format.space_after = Pt(4)
        group_run = group_paragraph.add_run("Group foundation routing:   ")
        group_run.font.name = FONT_NAME
        group_run.bold = True
        group_run.font.size = Pt(9.5)
        group_run.font.color.rgb = AMBER
        add_marked_run(
            group_paragraph,
            f"{group_foundation.get('foundation_name', 'a separate group foundation')} — "
            f"{group_foundation.get('explanation', '')}",
            base_size=9.5,
        )
        if group_foundation.get("source"):
            add_source_reference_paragraph(doc, group_foundation["source"], source_lookup)

    add_section_heading(doc, "Contact pathway")
    contact_pathway = analysis.get("contact_pathway", {}) if has_analysis else {}
    if contact_pathway.get("channel"):
        add_marked_paragraph(doc, contact_pathway["channel"], base_size=9.5)
        if contact_pathway.get("evidence"):
            add_evidence_paragraph(doc, contact_pathway["evidence"], source_lookup.get(contact_pathway.get("source", ""), ""))
    else:
        note_run = doc.add_paragraph().add_run(
            "No concrete outreach channel was identified in the fetched sources.")
        note_run.font.name = FONT_NAME
        style_small(note_run, 9, GREY)

    geographies = analysis.get("geographies", []) if has_analysis else []
    if geographies:
        add_section_heading(doc, "Geography")
        geo_paragraph = doc.add_paragraph()
        geo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        geo_paragraph.paragraph_format.space_after = Pt(4)
        label_run = geo_paragraph.add_run("Coverage:   ")
        label_run.font.name = FONT_NAME
        label_run.bold = True
        geo_text = ", ".join(g.get("place", "") for g in geographies)
        text_run = geo_paragraph.add_run(geo_text)
        text_run.font.name = FONT_NAME
        text_run.font.size = Pt(9.5)
        footnote_run = doc.add_paragraph().add_run("Shown for reference only — not used in scoring.")
        footnote_run.font.name = FONT_NAME
        style_small(footnote_run)

    add_section_heading(doc, "Source quality & CSR leadership note")
    if has_analysis and analysis.get("source_quality_assessment"):
        add_marked_paragraph(doc, analysis["source_quality_assessment"], base_size=9.5)
    if has_analysis and analysis.get("csr_head_note"):
        add_marked_paragraph(doc, analysis["csr_head_note"], base_size=9)
    if has_analysis and analysis.get("evidence_recency"):
        add_marked_paragraph(doc, analysis["evidence_recency"], base_size=9)

    add_section_heading(doc, "Sources consulted")
    for source in sources:
        if source.get("status") == "NOT_TRIED":
            continue
        source_paragraph = doc.add_paragraph()
        source_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        source_paragraph.paragraph_format.space_after = Pt(3)
        status_mark = "\u2714" if source.get("status") == "FOUND" else "\u2716"
        source_run = source_paragraph.add_run(
            f"{status_mark}   {SOURCE_LABELS.get(source.get('source_name', ''), source.get('source_name', ''))}    ")
        source_run.font.name = FONT_NAME
        source_run.font.size = Pt(9)
        source_run.font.color.rgb = TEAL_MID if source.get("status") == "FOUND" else GREY
        if source.get("url"):
            add_hyperlink(source_paragraph, source["url"], source["url"][:70])
        if source.get("source_name") == "people_search" and source.get("people_hits"):
            for hit in source["people_hits"][:10]:
                hit_paragraph = doc.add_paragraph()
                hit_paragraph.paragraph_format.left_indent = Cm(0.6)
                hit_paragraph.paragraph_format.space_after = Pt(2)
                hit_name = (hit.get("name") or "").strip()
                hit_title = (hit.get("title") or "").strip()
                hit_run = hit_paragraph.add_run(f"{hit_name}{' — ' + hit_title if hit_title else ''}   ")
                style_small(hit_run, 8.5, INK, italic=False)
                if hit.get("url"):
                    add_hyperlink(hit_paragraph, hit["url"], "LinkedIn", color="0A66C2", bold=True)

    if important_links:
        add_section_heading(doc, "Important links")
        for link in important_links:
            link_paragraph = doc.add_paragraph()
            link_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            link_paragraph.paragraph_format.space_after = Pt(3)
            add_hyperlink(link_paragraph, link.get("url", ""), link.get("label", link.get("url", "")))
            if link.get("relevance"):
                relevance_run = link_paragraph.add_run(f"   — {link['relevance']}")
                relevance_run.font.name = FONT_NAME
                relevance_run.font.size = Pt(8.5)

    footer_paragraph = doc.add_paragraph()
    footer_paragraph.paragraph_format.space_before = Pt(18)
    footer_run = footer_paragraph.add_run(
        "Generated by TAP CSR Research Agent  ·  fundraising@theapprenticeproject.org  ·  "
        "All data from public sources with citations  ·  verify every figure before outreach  ·  "
        "Internal use only")
    footer_run.font.name = FONT_NAME
    style_small(footer_run, 7.5)
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
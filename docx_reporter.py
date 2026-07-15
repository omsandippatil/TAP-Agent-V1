# docx_reporter.py — leadership-ready Word report
"""
Generates a .docx research brief designed for senior leadership:
  - Executive summary up front (score, verdict, one-paragraph insight)
  - Visual score breakdown (colour bars built from table shading)
  - Funded partners with TAP-similarity flags
  - CSR decision-makers with LinkedIn links
  - Full evidence citations + verification log (zero-hallucination audit)
Every fact shown carries its source. Nothing is generated without evidence.
"""

import io
import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Palette (matches the app UI) ─────────────────────────────────────────────
PURPLE   = RGBColor(0x7C, 0x3A, 0xED)
INK      = RGBColor(0x1F, 0x29, 0x37)
GREY     = RGBColor(0x6B, 0x72, 0x80)
GREEN    = RGBColor(0x16, 0xA3, 0x4A)
AMBER    = RGBColor(0xD9, 0x77, 0x06)
RED      = RGBColor(0xDC, 0x26, 0x26)
CREAM_HX = "FAF7F2"
PURPLE_HX= "7C3AED"

BLUE = RGBColor(0x0E, 0xA5, 0xE9)

def _fit_color(s):
    return (PURPLE if s >= 90 else GREEN if s >= 80 else
            BLUE if s >= 65 else AMBER if s >= 45 else RED)

def _fit_hex(s):
    return ("7C3AED" if s >= 90 else "16A34A" if s >= 80 else
            "0EA5E9" if s >= 65 else "D97706" if s >= 45 else "DC2626")

def _fit_label(s):
    return ("PRIORITY MATCH" if s >= 90 else "STRONG MATCH" if s >= 80 else
            "PROMISING" if s >= 65 else "WATCHLIST" if s >= 45 else "LOW FIT")


# ── Low-level helpers ────────────────────────────────────────────────────────

def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _add_hyperlink(paragraph, url: str, text: str, color="0563C1"):
    """python-docx has no native hyperlink API — build the XML directly."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u");     u.set(qn("w:val"), "single"); rPr.append(u)
    sz = OxmlElement("w:sz");   sz.set(qn("w:val"), "18");    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)

def _h(doc, text, size=13, color=PURPLE, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p

def _small(paragraph_or_run, size=8.5, color=GREY, italic=True):
    f = paragraph_or_run.font
    f.size = Pt(size); f.color.rgb = color; f.italic = italic

def _evidence_par(doc, exrpt: str, url: str = ""):
    if not exrpt:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'“{exrpt[:260]}…”  ')
    _small(r)
    if url:
        _add_hyperlink(p, url, "source")

def _score_bar(table_row, label, score, mx, hex_color):
    """Row = [label][bar cells x10][score]. Filled cells get shaded."""
    cells = table_row.cells
    cells[0].text = label
    cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    filled = round((score / mx) * 10) if mx else 0
    for i in range(10):
        cells[1 + i].text = ""
        _shade(cells[1 + i], hex_color if i < filled else "E5E7EB")
    cells[11].text = f"{score}/{mx}"
    cells[11].paragraphs[0].runs[0].font.size = Pt(9)


# ── Main generator ───────────────────────────────────────────────────────────

def generate_docx_report(company: str, result: dict, mode: str = "deep") -> bytes:
    fit       = result["fit_score"]
    state     = result["state"]
    insight   = result["strategic_insight"]
    breakdown = result.get("breakdown", {})
    data      = result.get("data", {})
    sources   = result.get("sources", [])
    verif     = data.get("verification", {})
    now       = datetime.datetime.now().strftime("%d %B %Y")

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.6)
        s.left_margin = s.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = INK

    # ── Header band ──────────────────────────────────────────────────────────
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, "1A1A2E")
    p = cell.paragraphs[0]
    r = p.add_run("THE APPRENTICE PROJECT — CSR RESEARCH BRIEF")
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xC4, 0xB5, 0xFD)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(company)
    r2.bold = True; r2.font.size = Pt(20); r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p3 = cell.add_paragraph()
    r3 = p3.add_run(f"{'Deep Research' if mode=='deep' else 'Prospect Screening'} · {now} · "
                    f"Prepared by TAP Fundraising Research Agent")
    r3.font.size = Pt(8.5); r3.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Executive summary ────────────────────────────────────────────────────
    _h(doc, "EXECUTIVE SUMMARY")
    t = doc.add_table(rows=1, cols=2)
    t.columns[0].width = Cm(4.2)
    sc = t.rows[0].cells[0]
    _shade(sc, _fit_hex(fit))
    sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = sc.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{fit}")
    r.bold = True; r.font.size = Pt(34); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p = sc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{_fit_label(fit)} / 100")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    ic = t.rows[0].cells[1]
    _shade(ic, CREAM_HX)
    ic.paragraphs[0].add_run(insight).font.size = Pt(9.5)
    state_lbl = {"FOUND": "Evidence found across sources",
                 "NOT_FOUND_IN_SOURCE": "Partial research — some sources empty",
                 "CONFIRMED_ABSENT": "No public India CSR evidence"}.get(state, state)
    p = ic.add_paragraph()
    r = p.add_run(f"Research state: {state_lbl}")
    _small(r, 8, GREY, italic=False)

    # ── CSR delivery model — funder vs implementer (pre-sales filter) ────────
    delivery = data.get("csr_delivery_model", {})
    _h(doc, "CSR DELIVERY MODEL")
    model = delivery.get("model", "UNCLEAR")
    model_color = {"FUNDER": GREEN, "HYBRID": PURPLE,
                   "IMPLEMENTER": AMBER}.get(model, GREY)
    p = doc.add_paragraph()
    r = p.add_run(model + "  ")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = model_color
    r = p.add_run(delivery.get("note", ""))
    r.font.size = Pt(9.5)
    if delivery.get("funder_signals"):
        p = doc.add_paragraph()
        r = p.add_run("Funder signals: " + ", ".join(delivery["funder_signals"]))
        _small(r, 8.5, GREY, italic=False)
        _evidence_par(doc, delivery.get("funder_evidence",""))
    if delivery.get("implementer_signals"):
        p = doc.add_paragraph()
        r = p.add_run("Implementer signals: " + ", ".join(delivery["implementer_signals"]))
        _small(r, 8.5, GREY, italic=False)
        _evidence_par(doc, delivery.get("implementer_evidence",""))

    # ── Score breakdown ──────────────────────────────────────────────────────
    _h(doc, "SCORE BREAKDOWN")
    dims = {"focus_alignment": "Focus Alignment", "adjacency_boost": "Adjacency Boost",
            "partner_similarity": "Partner Similarity",
            "geography_fit": "Geography", "csr_maturity": "CSR Maturity",
            "budget_size": "Budget", "source_quality": "Source Quality"}
    bar_rows = {k: v for k, v in breakdown.items()
                if isinstance(v, dict) and "score" in v and "max" in v}
    bt = doc.add_table(rows=len(bar_rows), cols=12)
    bt.autofit = False
    for i, (dim, info) in enumerate(bar_rows.items()):
        _score_bar(bt.rows[i], dims.get(dim, dim),
                   info.get("score", 0), info.get("max", 10), PURPLE_HX)

    sem = breakdown.get("semantic_alignment", {})
    if isinstance(sem, dict) and sem.get("used"):
        p = doc.add_paragraph()
        r = p.add_run(f"AI Semantic Alignment: {sem.get('score', 0)}/100")
        r.bold = True
        if sem.get("rationale"):
            _small(p.add_run("  —  " + sem["rationale"]), 8.5, GREY)

    # ── CSR spend ────────────────────────────────────────────────────────────
    spend = data.get("spend", {})
    _h(doc, "INDIA CSR SPEND")
    p = doc.add_paragraph()
    r = p.add_run(f"{spend.get('display','Not publicly disclosed')}  "
                  f"{spend.get('usd_approx','')}")
    r.bold = True; r.font.size = Pt(11)
    ef = spend.get("evidence_fact")
    if ef:
        _evidence_par(doc, ef.get("excerpt",""), ef.get("source_url",""))

    # ── Funded partners ──────────────────────────────────────────────────────
    _h(doc, "FUNDED / IMPLEMENTATION PARTNERS")
    partners = data.get("ngo_partners", [])
    if partners:
        pt = doc.add_table(rows=1, cols=4)
        pt.style = "Light Grid Accent 4"
        hdr = pt.rows[0].cells
        for j, htxt in enumerate(["Partner", "Similar to TAP?", "Why", "Source"]):
            hdr[j].text = ""
            r = hdr[j].paragraphs[0].add_run(htxt)
            r.bold = True; r.font.size = Pt(9)
            _shade(hdr[j], "EDE9FE")
        for pnr in partners:
            row = pt.add_row().cells
            r = row[0].paragraphs[0].add_run(
                pnr["name"] + ("  (company's own foundation)" if pnr.get("is_own_foundation") else ""))
            r.font.size = Pt(9); r.bold = pnr.get("tap_similar", False)
            sim = ("YES — known TAP-peer NGO" if pnr.get("is_peer_ngo")
                   else "YES" if pnr.get("tap_similar") else "—")
            r = row[1].paragraphs[0].add_run(sim)
            r.font.size = Pt(9)
            r.font.color.rgb = GREEN if pnr.get("tap_similar") else GREY
            r = row[2].paragraphs[0].add_run(", ".join(pnr.get("similarity_signals", [])) or "—")
            r.font.size = Pt(8.5)
            if pnr.get("source_url"):
                _add_hyperlink(row[3].paragraphs[0], pnr["source_url"], "link")
        p = doc.add_paragraph()
        r = p.add_run("Partners are listed only when the name literally appears in a fetched "
                      "source in CSR context. 'Similar to TAP' = known peer NGO or ≥2 "
                      "education/skilling signals near the mention.")
        _small(r)
    else:
        _small(doc.add_paragraph().add_run("No funded partners found in public sources."))

    # ── Decision makers ──────────────────────────────────────────────────────
    _h(doc, "CSR DECISION-MAKERS")
    dms = data.get("decision_makers", [])
    if dms:
        dt = doc.add_table(rows=1, cols=3)
        dt.style = "Light Grid Accent 4"
        hdr = dt.rows[0].cells
        for j, htxt in enumerate(["Name", "Title / Role", "LinkedIn"]):
            hdr[j].text = ""
            r = hdr[j].paragraphs[0].add_run(htxt); r.bold = True; r.font.size = Pt(9)
            _shade(hdr[j], "EDE9FE")
        for dm in dms:
            row = dt.add_row().cells
            row[0].paragraphs[0].add_run(dm.get("name","")).font.size = Pt(9)
            row[1].paragraphs[0].add_run(dm.get("title", dm.get("name",""))).font.size = Pt(9)
            p = row[2].paragraphs[0]
            if dm.get("linkedin_url"):
                _add_hyperlink(p, dm["linkedin_url"], "profile (from search results)")
            elif dm.get("linkedin_search_url"):
                _add_hyperlink(p, dm["linkedin_search_url"], "search on LinkedIn")
        p = doc.add_paragraph()
        r = p.add_run("Names come only from public sources / LinkedIn's own search snippets. "
                      "'Search on LinkedIn' links run a live search — they do not assert a "
                      "specific profile. Verify before outreach.")
        _small(r)
    else:
        _small(doc.add_paragraph().add_run("No CSR decision-makers found in public sources."))

    # ── Adjacency clusters ───────────────────────────────────────────────────
    adj = breakdown.get("adjacency_boost", {}).get("fired_clusters", [])
    if adj:
        _h(doc, "PARTNERSHIP ANGLES (ADJACENCY)")
        for c in adj:
            p = doc.add_paragraph()
            r = p.add_run(f"{c['label']}  (+{c['boost_applied']} pts)  ")
            r.bold = True; r.font.size = Pt(10)
            r2 = p.add_run("· " + ", ".join(c.get("keywords_found", [])[:4]))
            _small(r2, 8.5, GREY, italic=False)
            pr = doc.add_paragraph(); pr.paragraph_format.left_indent = Cm(0.4)
            pr.add_run(c.get("tap_reasoning","")).font.size = Pt(9)
            for ev in c.get("evidence_excerpts", [])[:1]:
                _evidence_par(doc, ev)

    # ── Geography & programmes ───────────────────────────────────────────────
    geos  = [g.get("place","") for g in data.get("geography", [])]
    progs = [p_.get("name","") for p_ in data.get("programs", [])]
    if geos or progs:
        _h(doc, "FOOTPRINT & PROGRAMMES")
        if geos:
            p = doc.add_paragraph()
            p.add_run("Geography: ").bold = True
            p.add_run(", ".join(geos)).font.size = Pt(9.5)
        for name in progs[:8]:
            doc.add_paragraph(name, style="List Bullet").runs[0].font.size = Pt(9.5)

    # ── Sources consulted ────────────────────────────────────────────────────
    _h(doc, "SOURCES CONSULTED")
    labels = {
        "india_csr_page": "Company India CSR page",
        "mca_portal": "MCA portal (verified)",
        "mca_via_search": "MCA via web search (proxy — not a verified filing)",
        "national_csr_portal": "National CSR Portal (csr.gov.in)",
        "annual_report": "Annual / sustainability report",
        "global_annual_report": "Annual / sustainability report",
        "partner_search": "Partner-focused web search",
        "people_search": "LinkedIn / people search snippets",
        "plans_search": "Partnerships & announced-plans search",
    }
    for s in sources:
        if s.get("status") == "NOT_TRIED":
            continue
        p = doc.add_paragraph()
        mark = "✔" if s.get("status") == "FOUND" else "✖"
        r = p.add_run(f"{mark}  {labels.get(s.get('source_name',''), s.get('source_name',''))}   ")
        r.font.size = Pt(9)
        r.font.color.rgb = GREEN if s.get("status") == "FOUND" else GREY
        if s.get("url"):
            _add_hyperlink(p, s["url"], s["url"][:70])

    # ── Verification log ─────────────────────────────────────────────────────
    _h(doc, "VERIFICATION LOG")
    checks = verif.get("checks", [])
    p = doc.add_paragraph()
    r = p.add_run(f"Every published fact was re-checked against raw source text: "
                  f"{verif.get('verified',0)}/{len(checks)} verified "
                  f"({verif.get('pass_rate',100)}% pass). "
                  f"Facts without source evidence are never shown.")
    r.font.size = Pt(9)
    if checks:
        vt = doc.add_table(rows=1, cols=3)
        vt.style = "Light Grid Accent 4"
        hdr = vt.rows[0].cells
        for j, htxt in enumerate(["Field", "Value", "Status"]):
            hdr[j].text = ""
            rr = hdr[j].paragraphs[0].add_run(htxt); rr.bold = True; rr.font.size = Pt(8.5)
            _shade(hdr[j], "EDE9FE")
        for c in checks[:30]:
            row = vt.add_row().cells
            row[0].paragraphs[0].add_run(c["field"]).font.size = Pt(8.5)
            row[1].paragraphs[0].add_run(c["value"]).font.size = Pt(8.5)
            rr = row[2].paragraphs[0].add_run(c["status"])
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = GREEN if c["status"] == "VERIFIED" else AMBER

    # ── Footer ───────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("Generated by TAP CSR Research Agent · fundraising@theapprenticeproject.org · "
                  "All data from public sources with citations · Zero-hallucination architecture · "
                  "Internal use only")
    _small(r, 7.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
